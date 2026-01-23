from mcp.server.fastmcp import FastMCP
from fastapi import FastAPI, HTTPException
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
import httpx
import os
import asyncio
import io
import re
import fitz  # PyMuPDF for PDF extraction

mcp = FastMCP("docai-mcp")

BACKEND_URL = os.getenv("BACKEND_URL", "http://backend:8000/api/v1")
AI_API_KEY = os.getenv("AI_API_KEY")
AI_API_BASE_URL = os.getenv("AI_API_BASE_URL", "https://integrate.api.nvidia.com/v1/chat/completions")
AI_MODEL_NAME = os.getenv("AI_MODEL_NAME", "glm4.7")

class AIClient:
    def __init__(self):
        self.api_key = (AI_API_KEY or "").strip()
        self.api_url = (AI_API_BASE_URL or "").strip()
        self.model = AI_MODEL_NAME

    def _resolve_url(self) -> str:
        if not self.api_url:
            return ""
        if self.api_url.endswith("/chat/completions"):
            return self.api_url
        return self.api_url.rstrip("/") + "/chat/completions"

    def _resolve_model(self, override_model: str | None) -> str:
        model = (override_model or self.model or "").strip()
        return model or "glm4.7"
        
    def generate_completion(self, prompt: str, model: str | None = None):
        url = self._resolve_url()
        if not self.api_key or not url:
            return f"Mock AI Response (Key Missing): {prompt[:50]}..."

        selected_model = self._resolve_model(model)
        try:
            with httpx.Client(timeout=60.0) as client:
                resp = client.post(
                    url,
                    headers={"Authorization": f"Bearer {self.api_key}"},
                    json={
                        "model": selected_model,
                        "messages": [{"role": "user", "content": prompt}],
                        "stream": False,
                    },
                )
                if resp.status_code >= 400:
                    return f"AI Error: {resp.status_code}: {resp.text}"
                data = resp.json()
                return data["choices"][0]["message"]["content"]
        except Exception as e:
            print(f"AI Error: {e}")
            return f"AI Error: {str(e)}"

ai_client = AIClient()

def _env_int(name: str, default: int) -> int:
    try:
        val = int((os.getenv(name) or "").strip())
        return val if val > 0 else default
    except Exception:
        return default

def _normalize_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()

def _split_text_into_chunks(text: str, max_chars: int, overlap_chars: int, max_chunks: int) -> list[str]:
    text = _normalize_text(text)
    if not text:
        return []
    if len(text) <= max_chars:
        return [text]

    paragraphs = [p for p in re.split(r"\n{2,}", text) if p.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    def flush():
        nonlocal current, current_len
        if not current:
            return
        chunk = "\n\n".join(current).strip()
        if chunk:
            chunks.append(chunk)
        current = []
        current_len = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(para) > max_chars:
            flush()
            start = 0
            while start < len(para) and len(chunks) < max_chunks:
                end = min(start + max_chars, len(para))
                chunks.append(para[start:end].strip())
                if end >= len(para):
                    break
                start = max(0, end - overlap_chars)
            continue

        if current_len + len(para) + (2 if current else 0) > max_chars:
            flush()
        current.append(para)
        current_len += len(para) + (2 if current_len else 0)
        if len(chunks) >= max_chunks:
            break

    flush()
    if len(chunks) > max_chunks:
        chunks = chunks[:max_chunks]
    return chunks

def _strip_think(text: str) -> str:
    return re.sub(r"<think>[\s\S]*?</think>", "", text or "").strip()

def _clean_json_like(text: str) -> str:
    text = (text or "").strip()
    text = re.sub(r",\s*([}\]])", r"\1", text)
    return text.strip()

def _balanced_json_substrings(text: str) -> list[str]:
    text = text or ""
    candidates: list[str] = []
    for start_char, end_char in [("{", "}"), ("[", "]")]:
        start_positions = [m.start() for m in re.finditer(re.escape(start_char), text)]
        for start in start_positions[:30]:
            depth = 0
            in_string = False
            escape = False
            for i in range(start, len(text)):
                c = text[i]
                if in_string:
                    if escape:
                        escape = False
                    elif c == "\\":
                        escape = True
                    elif c == '"':
                        in_string = False
                    continue
                if c == '"':
                    in_string = True
                    continue
                if c == start_char:
                    depth += 1
                elif c == end_char:
                    depth -= 1
                    if depth == 0:
                        frag = text[start : i + 1].strip()
                        if frag:
                            candidates.append(frag)
                        break
    return candidates

def _extract_json_candidates(text: str) -> list[str]:
    text = _strip_think(text)
    candidates: list[str] = []

    fenced = re.findall(r"```(?:json)?\s*([\s\S]*?)\s*```", text, flags=re.IGNORECASE)
    for block in fenced:
        block = block.strip()
        if block:
            candidates.append(block)

    candidates.extend(_balanced_json_substrings(text))

    raw = text.strip()
    if raw and raw[0] in "{[":
        candidates.append(raw)
    return candidates

def _robust_json_loads(text: str):
    last_err: Exception | None = None
    for cand in _extract_json_candidates(text):
        cand = _clean_json_like(cand)
        if not cand:
            continue
        try:
            return json.loads(cand)
        except Exception as e:
            last_err = e
            continue
    if last_err:
        raise last_err
    raise json.JSONDecodeError("No JSON candidate found", text or "", 0)

def _deep_merge(a, b):
    if a is None:
        return b
    if b is None:
        return a
    if isinstance(a, dict) and isinstance(b, dict):
        out = dict(a)
        for k, v in b.items():
            out[k] = _deep_merge(out.get(k), v)
        return out
    if isinstance(a, list) and isinstance(b, list):
        out = list(a)
        for item in b:
            if item not in out:
                out.append(item)
        return out
    if isinstance(a, str) and isinstance(b, str):
        a2 = a.strip()
        b2 = b.strip()
        if not a2:
            return b
        if not b2:
            return a
        return b if len(b2) >= len(a2) else a
    return b

def _hierarchical_summarize(text: str, ai_model: str | None, target_chars: int) -> str:
    text = _normalize_text(text)
    if not text:
        return ""
    if len(text) <= target_chars:
        return text

    chunk_size = _env_int("AI_CHUNK_SIZE_CHARS", 8000)
    overlap = _env_int("AI_CHUNK_OVERLAP_CHARS", 400)
    max_chunks = _env_int("AI_MAX_CHUNKS", 8)

    chunks = _split_text_into_chunks(text, chunk_size, overlap, max_chunks)
    summaries: list[str] = []
    for idx, chunk in enumerate(chunks, start=1):
        prompt = (
            "你是信息抽取助手。请仅基于这段内容，提炼可用于后续结构化抽取的关键信息，要求：\n"
            "1) 输出为要点列表，每条不超过 80 字；2) 保留关键实体、时间、金额、条款编号、层级标题；\n"
            "3) 不要编造；4) 不要输出除要点外的任何解释。\n\n"
            f"【片段 {idx}/{len(chunks)}】\n{chunk}"
        )
        s = ai_client.generate_completion(prompt, model=ai_model)
        s = _normalize_text(_strip_think(s))
        if s:
            summaries.append(s)

    merged = _normalize_text("\n".join(summaries))
    if len(merged) <= target_chars:
        return merged

    prompt2 = (
        "请将以下要点再压缩为更短的要点列表，要求：\n"
        "1) 保留字段抽取所需的实体、数字、条款编号、章节层级；2) 不要编造；3) 只输出要点。\n\n"
        f"{merged}"
    )
    s2 = ai_client.generate_completion(prompt2, model=ai_model)
    return _normalize_text(_strip_think(s2))

def _docx_iter_block_items(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)

def _docx_paragraph_to_markdown(para: Paragraph) -> str:
    text = (para.text or "").strip()
    if not text:
        return ""
    style_name = (getattr(getattr(para, "style", None), "name", None) or "").strip().lower()
    if style_name in {"title", "标题"}:
        return f"# {text}"
    m = re.match(r"heading\s*(\d+)", style_name)
    if m:
        level = int(m.group(1))
        level = min(max(level, 1), 6)
        return f"{'#' * level} {text}"
    return text

def _docx_table_to_markdown(table: Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        row_vals: list[str] = []
        for cell in row.cells:
            cell_text = _normalize_text(cell.text)
            row_vals.append(cell_text.replace("\n", "<br>") if cell_text else "")
        rows.append(row_vals)
    if not rows:
        return ""

    max_cols = max((len(r) for r in rows), default=0)
    if max_cols <= 0:
        return ""
    for r in rows:
        if len(r) < max_cols:
            r.extend([""] * (max_cols - len(r)))

    header = rows[0]
    has_header = any(v.strip() for v in header)
    if not has_header:
        header = [f"列{i+1}" for i in range(max_cols)]
        data_rows = rows
    else:
        data_rows = rows[1:] if len(rows) > 1 else []

    def fmt_row(r: list[str]) -> str:
        return "| " + " | ".join([v.replace("|", "\\|") for v in r]) + " |"

    out = [fmt_row(header), "| " + " | ".join(["---"] * max_cols) + " |"]
    for r in data_rows:
        out.append(fmt_row(r))
    return "\n".join(out)

def _extract_pdf_text(file_content: bytes) -> tuple[str, dict]:
    pdf_doc = fitz.open(stream=file_content, filetype="pdf")
    pages: list[str] = []
    for i, page in enumerate(pdf_doc, start=1):
        blocks = page.get_text("blocks") or []
        blocks_sorted = sorted(blocks, key=lambda b: (b[1], b[0]))
        parts = []
        for b in blocks_sorted:
            t = (b[4] or "").strip()
            if t:
                parts.append(t)
        text = _normalize_text("\n".join(parts)) if parts else _normalize_text(page.get_text() or "")
        if text:
            pages.append(f"--- Page {i} ---\n{text}")
    meta = {"pages": pdf_doc.page_count}
    pdf_doc.close()
    return (_normalize_text("\n\n".join(pages)), meta)

def _extract_docx_text(file_content: bytes) -> tuple[str, dict]:
    doc = Document(io.BytesIO(file_content))
    parts: list[str] = []
    para_count = 0
    table_count = 0
    for block in _docx_iter_block_items(doc):
        if isinstance(block, Paragraph):
            md = _docx_paragraph_to_markdown(block)
            if md:
                parts.append(md)
                para_count += 1
        elif isinstance(block, Table):
            md = _docx_table_to_markdown(block)
            if md:
                parts.append(md)
                table_count += 1
    return (_normalize_text("\n\n".join(parts)), {"paragraphs": para_count, "tables": table_count})

def _infer_template_type(text: str) -> str:
    t = (text or "").lower()
    score = {
        "resume": 0,
        "report": 0,
        "meeting": 0,
        "contract": 0,
        "proposal": 0,
        "invoice": 0,
    }
    for kw in ["简历", "教育背景", "工作经历", "自我评价", "技能"]:
        if kw in t:
            score["resume"] += 1
    for kw in ["项目概述", "项目目标", "实施过程", "主要成果", "总结", "报告"]:
        if kw in t:
            score["report"] += 1
    for kw in ["会议纪要", "参会", "会议时间", "议题", "决议", "行动项"]:
        if kw in t:
            score["meeting"] += 1
    for kw in ["合同", "甲方", "乙方", "违约", "争议解决", "付款"]:
        if kw in t:
            score["contract"] += 1
    for kw in ["提案", "项目背景", "实施方案", "预算", "风险评估", "时间计划"]:
        if kw in t:
            score["proposal"] += 1
    for kw in ["发票", "金额", "税", "开票", "商品/服务"]:
        if kw in t:
            score["invoice"] += 1
    best = max(score.items(), key=lambda x: x[1])
    return best[0] if best[1] > 0 else "default"

async def download_file_from_backend(file_id: str) -> bytes:
    url = f"{BACKEND_URL}/files/{file_id}/download"
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.get(url)
            if resp.status_code == 200:
                return resp.content
            raise Exception(f"Status {resp.status_code}: {resp.text[:200]}")
    except Exception as e:
        raise Exception(f"Download error: {str(e)}")

async def upload_generated_document(file_content: bytes, filename: str) -> str:
    """上传生成的文档到后端，失败时抛出异常"""
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            files = {'file': (filename, io.BytesIO(file_content), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            resp = await client.post(f"{BACKEND_URL}/files/upload", files=files)
            
            if resp.status_code == 200:
                file_id = resp.json().get('fileId')
                if file_id:
                    print(f"Successfully uploaded document: {filename} -> {file_id}")
                    return file_id
                else:
                    raise Exception(f"Upload succeeded but no fileId returned: {resp.json()}")
            else:
                raise Exception(f"Upload failed with status {resp.status_code}: {resp.text}")
    except httpx.TimeoutException:
        raise Exception(f"Upload timeout for {filename}")
    except httpx.RequestError as e:
        raise Exception(f"Network error uploading {filename}: {str(e)}")
    except Exception as e:
        print(f"Failed to upload generated document: {e}")
        raise

def set_chinese_font(run, font_name='微软雅黑', size=12, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def _format_structured_text(val) -> str:
    if val is None:
        return ""
    if isinstance(val, str):
        return val.strip()
    if isinstance(val, (int, float, bool)):
        return str(val)
    if isinstance(val, dict):
        parts: list[str] = []
        for k, v in val.items():
            sv = _format_structured_text(v)
            if sv.strip():
                parts.append(f"{k}：{sv}")
        if parts:
            return "\n".join(parts).strip()
        try:
            return json.dumps(val, ensure_ascii=False)
        except Exception:
            return str(val)
    if isinstance(val, list):
        if not val:
            return ""
        if all(isinstance(x, dict) for x in val):
            lines: list[str] = []
            for item in val:
                if not isinstance(item, dict):
                    continue
                mapped = {str(k): _format_structured_text(v) for k, v in item.items()}
                key_map = {
                    "item": "事项",
                    "deadline": "截止",
                    "owner": "负责人",
                    "assignee": "负责人",
                    "date": "日期",
                    "time": "时间",
                }
                if "item" in mapped or "事项" in mapped:
                    text = (mapped.get("事项") or mapped.get("item") or "").strip()
                    owner = (mapped.get("负责人") or mapped.get("owner") or mapped.get("assignee") or "").strip()
                    deadline = (mapped.get("截止") or mapped.get("deadline") or "").strip()
                    extra: list[str] = []
                    if owner:
                        extra.append(f"负责人：{owner}")
                    if deadline:
                        extra.append(f"截止：{deadline}")
                    line = text
                    if extra:
                        line = (line + "（" + "；".join(extra) + "）").strip()
                    if line:
                        lines.append(f"• {line}")
                else:
                    kv: list[str] = []
                    for k, v in mapped.items():
                        kk = key_map.get(k, k)
                        vv = (v or "").strip()
                        if vv:
                            kv.append(f"{kk}：{vv}")
                    if kv:
                        lines.append("• " + "；".join(kv))
            return "\n".join(lines).strip()
        if all(isinstance(x, str) for x in val):
            items = [f"• {x.strip()}" for x in val if isinstance(x, str) and x.strip()]
            return "\n".join(items).strip()
        items2 = []
        for x in val:
            sx = _format_structured_text(x)
            if sx.strip():
                items2.append(f"• {sx.strip()}")
        return "\n".join(items2).strip()
    try:
        return json.dumps(val, ensure_ascii=False)
    except Exception:
        return str(val)

def safe_str(val):
    """安全转换为字符串"""
    return _format_structured_text(val)

def parse_ai_content_for_template(content: str, template_type: str, ai_model: str | None = None) -> dict:
    """使用AI解析内容并提取结构化数据用于填充模板"""
    template_prompts = {
        'resume': """请从以下内容中提取简历信息，以JSON格式返回：
{
    "title": "个人简历",
    "personal_info": {"姓名": "", "电话": "", "邮箱": "", "地址": ""},
    "education": [{"学校": "", "专业": "", "学历": "", "时间": ""}],
    "work_experience": [{"公司": "", "职位": "", "时间": "", "职责": ""}],
    "skills": ["技能1", "技能2"],
    "self_evaluation": "自我评价内容"
}
如果某些信息不存在，请根据内容合理推断或留空。

原始内容：
""",
        'report': """请从以下内容中提取报告信息，以JSON格式返回：
{
    "title": "报告标题",
    "sections": {
        "项目概述": "内容",
        "项目目标": "内容",
        "实施过程": "内容",
        "主要成果": "内容",
        "存在问题": "内容",
        "改进建议": "内容",
        "总结": "内容"
    }
}
请根据原始内容填充各部分，如内容不足可合理扩展。

原始内容：
""",
        'meeting': """请从以下内容中提取会议纪要信息，以JSON格式返回：
{
    "title": "会议纪要",
    "meeting_time": "会议时间",
    "meeting_place": "会议地点",
    "attendees": "参会人员",
    "topics": "会议议题",
    "discussion": "讨论内容",
    "decisions": "决议事项",
    "actions": "后续行动"
}
请根据原始内容填充，如内容不足可合理推断。

原始内容：
""",
        'contract': """请从以下内容中提取合同信息，以JSON格式返回：
{
    "title": "合同协议",
    "party_a": {"名称": "", "地址": "", "联系人": "", "电话": ""},
    "party_b": {"名称": "", "地址": "", "联系人": "", "电话": ""},
    "content": "合同具体条款",
    "period": {"开始时间": "", "结束时间": ""},
    "payment": "付款条款",
    "breach": "违约责任",
    "dispute": "争议解决方式"
}
请根据原始内容填充。

原始内容：
""",
        'proposal': """请从以下内容中提取项目提案信息，以JSON格式返回：
{
    "title": "项目提案标题",
    "sections": {
        "项目背景": "内容",
        "项目目标": "内容",
        "项目内容": "内容",
        "实施方案": "内容",
        "时间计划": "内容",
        "预算说明": "内容",
        "预期成果": "内容",
        "风险评估": "内容"
    }
}
请根据原始内容填充。

原始内容：
"""
    }
    
    prompt = template_prompts.get(template_type, "")
    if not prompt:
        return {"raw_content": content}
    
    strict_suffix = "\n\n要求：只输出严格的 JSON（不要代码块、不要说明文字）。"

    max_direct = _env_int("AI_MAX_DIRECT_EXTRACT_CHARS", 12000)
    content_norm = _normalize_text(content)
    use_chunked = len(content_norm) > max_direct

    if use_chunked:
        extract_chunk_size = _env_int("AI_EXTRACT_CHUNK_SIZE_CHARS", 9000)
        overlap = _env_int("AI_EXTRACT_CHUNK_OVERLAP_CHARS", 400)
        max_chunks = _env_int("AI_EXTRACT_MAX_CHUNKS", _env_int("AI_MAX_CHUNKS", 8))
        chunks = _split_text_into_chunks(content_norm, extract_chunk_size, overlap, max_chunks)

        merged: dict = {}
        parsed_any = False
        for idx, chunk in enumerate(chunks, start=1):
            chunk_prompt = (
                prompt
                + chunk
                + "\n\n补充要求：仅基于该片段抽取，不要推断不存在的信息；缺失字段留空。"
                + strict_suffix
                + f"\n\n【片段 {idx}/{len(chunks)}】"
            )
            ai_resp = _strip_think(ai_client.generate_completion(chunk_prompt, model=ai_model))
            try:
                partial = _robust_json_loads(ai_resp)
                if isinstance(partial, dict):
                    merged = _deep_merge(merged, partial)
                    parsed_any = True
            except Exception:
                continue

        if parsed_any:
            merged["__source_length"] = len(content_norm)
            merged["__chunked"] = True
            merged["__chunks"] = len(chunks)
            return merged

        target_chars = _env_int("AI_SUMMARY_TARGET_CHARS", 12000)
        content_for_ai = _hierarchical_summarize(content_norm, ai_model=ai_model, target_chars=target_chars)
    else:
        content_for_ai = content_norm

    full_prompt = prompt + content_for_ai + strict_suffix
    ai_response = _strip_think(ai_client.generate_completion(full_prompt, model=ai_model))

    try:
        data = _robust_json_loads(ai_response)
        if isinstance(data, dict):
            data["__source_length"] = len(content or "")
        return data
    except Exception as e:
        repair_prompt = (
            "请把下面内容修正为严格 JSON，必须与用户要求的 JSON 结构保持一致，且只输出 JSON。\n\n"
            f"{ai_response}"
        )
        repaired = ai_client.generate_completion(repair_prompt, model=ai_model)
        repaired = _strip_think(repaired)
        try:
            data = _robust_json_loads(repaired)
            if isinstance(data, dict):
                data["__source_length"] = len(content or "")
                data["__repaired"] = True
            return data
        except Exception:
            excerpt_head = (content or "")[:20000]
            excerpt_tail = (content or "")[-5000:] if (content and len(content) > 25000) else ""
            return {
                "raw_content_excerpt": _normalize_text(excerpt_head + ("\n\n...[TRUNCATED]...\n\n" + excerpt_tail if excerpt_tail else "")),
                "title": "文档",
                "sections": {"内容": _normalize_text(excerpt_head)},
                "__parse_error": str(e),
                "__source_length": len(content or "")
            }

def create_document_from_content(content: str, template_type: str = 'default', ai_model: str | None = None) -> bytes:
    """根据内容和模板类型创建文档，使用AI解析内容填充模板"""
    doc = Document()
    
    # 如果有实际内容，使用AI解析并填充
    if content and content.strip() and template_type != 'default':
        parsed_data = parse_ai_content_for_template(content, template_type, ai_model=ai_model)
        
        if template_type == 'resume':
            title = parsed_data.get('title', '个人简历')
            doc.add_heading(safe_str(title), 0)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=18, bold=True)
            
            # 个人信息
            doc.add_heading('个人信息', level=1)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
            personal_info = parsed_data.get('personal_info', {})
            if isinstance(personal_info, dict):
                for key in ['姓名', '电话', '邮箱', '地址']:
                    p = doc.add_paragraph()
                    p.add_run(f'{key}：')
                    set_chinese_font(p.runs[-1], bold=True)
                    value = personal_info.get(key, '') or ''
                    p.add_run(safe_str(value) if value else '[待填写]')
                    set_chinese_font(p.runs[-1])
            doc.add_paragraph()
            
            # 教育背景
            doc.add_heading('教育背景', level=1)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
            education_list = parsed_data.get('education', [])
            if education_list and isinstance(education_list, list):
                for edu in education_list:
                    if isinstance(edu, dict):
                        for key in ['学校', '专业', '学历', '时间']:
                            p = doc.add_paragraph()
                            p.add_run(f'{key}：')
                            set_chinese_font(p.runs[-1], bold=True)
                            value = edu.get(key, '') or ''
                            p.add_run(safe_str(value) if value else '[待填写]')
                            set_chinese_font(p.runs[-1])
                        doc.add_paragraph()
            else:
                p = doc.add_paragraph('[待填写教育背景]')
                set_chinese_font(p.runs[-1])
            
            # 工作经历
            doc.add_heading('工作经历', level=1)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
            work_list = parsed_data.get('work_experience', [])
            if work_list and isinstance(work_list, list):
                for work in work_list:
                    if isinstance(work, dict):
                        for key in ['公司', '职位', '时间', '职责']:
                            p = doc.add_paragraph()
                            p.add_run(f'{key}：')
                            set_chinese_font(p.runs[-1], bold=True)
                            value = work.get(key, '') or ''
                            p.add_run(safe_str(value) if value else '[待填写]')
                            set_chinese_font(p.runs[-1])
                        doc.add_paragraph()
            else:
                p = doc.add_paragraph('[待填写工作经历]')
                set_chinese_font(p.runs[-1])
            
            # 技能特长
            doc.add_heading('技能特长', level=1)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
            skills = parsed_data.get('skills', [])
            if skills and isinstance(skills, list):
                p = doc.add_paragraph('• ' + '\n• '.join([safe_str(s) for s in skills]))
                set_chinese_font(p.runs[-1])
            else:
                p = doc.add_paragraph('[待填写技能]')
                set_chinese_font(p.runs[-1])
            doc.add_paragraph()
            
            # 自我评价
            doc.add_heading('自我评价', level=1)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
            self_eval = parsed_data.get('self_evaluation', '')
            p = doc.add_paragraph(safe_str(self_eval) if self_eval else '[待填写自我评价]')
            set_chinese_font(p.runs[-1])
        
        elif template_type == 'report':
            title = parsed_data.get('title', '项目报告')
            doc.add_heading(safe_str(title), 0)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=18, bold=True)
            
            sections = parsed_data.get('sections', {})
            section_order = ['项目概述', '项目目标', '实施过程', '主要成果', '存在问题', '改进建议', '总结']
            
            if isinstance(sections, dict):
                for section_name in section_order:
                    doc.add_heading(section_name, level=1)
                    set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
                    section_content = sections.get(section_name, '')
                    p = doc.add_paragraph(safe_str(section_content) if section_content else '[待填写内容]')
                    set_chinese_font(p.runs[-1])
        
        elif template_type == 'meeting':
            title = parsed_data.get('title', '会议纪要')
            doc.add_heading(safe_str(title), 0)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=18, bold=True)
            
            # 基本信息
            for key, field in [('会议时间', 'meeting_time'), ('会议地点', 'meeting_place'), ('参会人员', 'attendees')]:
                p = doc.add_paragraph()
                p.add_run(f'{key}：')
                set_chinese_font(p.runs[-1], bold=True)
                value = parsed_data.get(field, '')
                p.add_run(safe_str(value) if value else '[待填写]')
                set_chinese_font(p.runs[-1])
            
            # 各部分内容
            for section_name, field in [('会议议题', 'topics'), ('讨论内容', 'discussion'), ('决议事项', 'decisions'), ('后续行动', 'actions')]:
                doc.add_heading(section_name, level=1)
                set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
                section_content = parsed_data.get(field, '')
                p = doc.add_paragraph(safe_str(section_content) if section_content else '[待填写]')
                set_chinese_font(p.runs[-1])
        
        elif template_type == 'contract':
            title = parsed_data.get('title', '合同协议')
            doc.add_heading(safe_str(title), 0)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=18, bold=True)
            
            # 甲方信息
            doc.add_heading('甲方信息', level=1)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
            party_a = parsed_data.get('party_a', {})
            if isinstance(party_a, dict):
                for key in ['名称', '地址', '联系人', '电话']:
                    p = doc.add_paragraph()
                    p.add_run(f'{key}：')
                    set_chinese_font(p.runs[-1], bold=True)
                    value = party_a.get(key, '')
                    p.add_run(safe_str(value) if value else '[待填写]')
                    set_chinese_font(p.runs[-1])
            doc.add_paragraph()
            
            # 乙方信息
            doc.add_heading('乙方信息', level=1)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
            party_b = parsed_data.get('party_b', {})
            if isinstance(party_b, dict):
                for key in ['名称', '地址', '联系人', '电话']:
                    p = doc.add_paragraph()
                    p.add_run(f'{key}：')
                    set_chinese_font(p.runs[-1], bold=True)
                    value = party_b.get(key, '')
                    p.add_run(safe_str(value) if value else '[待填写]')
                    set_chinese_font(p.runs[-1])
            doc.add_paragraph()
            
            # 其他部分
            for section_name, field in [('合同内容', 'content'), ('付款方式', 'payment'), ('违约责任', 'breach'), ('争议解决', 'dispute')]:
                doc.add_heading(section_name, level=1)
                set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
                section_content = parsed_data.get(field, '')
                p = doc.add_paragraph(safe_str(section_content) if section_content else '[待填写]')
                set_chinese_font(p.runs[-1])
            
            # 合同期限
            doc.add_heading('合同期限', level=1)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
            period = parsed_data.get('period', {})
            if isinstance(period, dict):
                for key in ['开始时间', '结束时间']:
                    p = doc.add_paragraph()
                    p.add_run(f'{key}：')
                    set_chinese_font(p.runs[-1], bold=True)
                    value = period.get(key, '')
                    p.add_run(safe_str(value) if value else '[待填写]')
                    set_chinese_font(p.runs[-1])
        
        elif template_type == 'proposal':
            title = parsed_data.get('title', '项目提案')
            doc.add_heading(safe_str(title), 0)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=18, bold=True)
            
            sections = parsed_data.get('sections', {})
            section_order = ['项目背景', '项目目标', '项目内容', '实施方案', '时间计划', '预算说明', '预期成果', '风险评估']
            
            if isinstance(sections, dict):
                for section_name in section_order:
                    doc.add_heading(section_name, level=1)
                    set_chinese_font(doc.paragraphs[-1].runs[0], size=14, bold=True)
                    section_content = sections.get(section_name, '')
                    p = doc.add_paragraph(safe_str(section_content) if section_content else '[待填写内容]')
                    set_chinese_font(p.runs[-1])
        
        else:
            # 如果有raw_content，直接输出
            raw = parsed_data.get('raw_content', content)
            doc.add_heading('文档', 0)
            set_chinese_font(doc.paragraphs[-1].runs[0], size=18, bold=True)
            p = doc.add_paragraph(safe_str(raw))
            set_chinese_font(p.runs[-1])
    
    elif template_type == 'invoice':
        # 发票模板保持原样，需要用户手动填写
        doc.add_heading('发票', 0)
        set_chinese_font(doc.paragraphs[-1].runs[0], size=18, bold=True)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['商品/服务名称', '数量', '单价', '金额']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_chinese_font(hdr_cells[i].paragraphs[0].runs[0], bold=True)
        
        for _ in range(3):
            row_cells = table.add_row().cells
            for i in range(4):
                row_cells[i].text = '[待填写]'
                set_chinese_font(row_cells[i].paragraphs[0].runs[0])
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('合计金额：')
        set_chinese_font(p.runs[-1], bold=True)
        p.add_run('[待填写]')
        
        p = doc.add_paragraph()
        p.add_run('开票日期：')
        set_chinese_font(p.runs[-1], bold=True)
        p.add_run('[待填写]')
    
    else:
        # default类型：直接输出内容
        doc.add_heading('文档', 0)
        set_chinese_font(doc.paragraphs[-1].runs[0], size=18, bold=True)
        
        if content and content.strip():
            # 将内容按段落分割并添加
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # 检查是否是标题行
                    if para_text.startswith('#'):
                        # Markdown标题处理
                        level = len(para_text.split()[0].replace('#', '')) if para_text.startswith('#') else 1
                        level = min(level, 3)
                        title_text = para_text.lstrip('#').strip()
                        doc.add_heading(title_text, level=level)
                        if doc.paragraphs[-1].runs:
                            set_chinese_font(doc.paragraphs[-1].runs[0], size=14 - level, bold=True)
                    else:
                        p = doc.add_paragraph(para_text.strip())
                        set_chinese_font(p.runs[-1])
        else:
            p = doc.add_paragraph('[待填写内容]')
            set_chinese_font(p.runs[-1])
    
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def modify_document_with_content(original_content: bytes, modifications: str, ai_model: str | None = None) -> bytes:
    try:
        doc = Document(io.BytesIO(original_content))
        
        prompt = f"""
        根据以下修改要求，对文档进行修改。请以结构化的方式返回修改指令：
        
        修改要求：{modifications}
        
        请以JSON格式返回修改内容，格式如下：
        {{
            "title": "新标题（如需修改）",
            "add_paragraphs": ["段落1内容", "段落2内容"],
            "replace_sections": {{"原文本": "新文本"}},
            "format_changes": {{"font_size": 12, "bold": true}}
        }}
        """
        
        ai_response = ai_client.generate_completion(prompt, model=ai_model)
        
        try:
            if "```json" in ai_response:
                ai_response = ai_response.split("```json")[1].split("```")[0].strip()
            elif "```" in ai_response:
                ai_response = ai_response.split("```")[1].split("```")[0].strip()
            
            changes = json.loads(ai_response)
            
            if changes.get("title"):
                if doc.paragraphs:
                    doc.paragraphs[0].text = changes["title"]
                    set_chinese_font(doc.paragraphs[0].runs[0], size=18, bold=True)
            
            for para_text in changes.get("add_paragraphs", []):
                p = doc.add_paragraph(para_text)
                set_chinese_font(p.runs[-1])
            
            for old_text, new_text in changes.get("replace_sections", {}).items():
                for para in doc.paragraphs:
                    if old_text in para.text:
                        para.text = para.text.replace(old_text, new_text)
                        set_chinese_font(para.runs[0])
            
        except json.JSONDecodeError:
            p = doc.add_paragraph()
            p.add_run("AI 修改建议：")
            set_chinese_font(p.runs[-1], bold=True)
            p.add_run(ai_response)
            set_chinese_font(p.runs[-1])
        
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
        
    except Exception as e:
        print(f"Error modifying document: {e}")
        return original_content

async def _analyze_document_logic(file_id: str, analysis_type: str, ai_model: str | None = None):
    file_content = await download_file_from_backend(file_id)
    
    if not file_content:
        return {"error": "Failed to download document"}

    extracted_text = ""
    meta: dict = {}
    try:
        extracted_text, meta = _extract_pdf_text(file_content)
        meta["detected_type"] = "pdf"
    except Exception:
        try:
            extracted_text, meta = _extract_docx_text(file_content)
            meta["detected_type"] = "docx"
        except Exception:
            try:
                extracted_text = file_content.decode("utf-8", errors="ignore")
                meta["detected_type"] = "text"
            except Exception:
                extracted_text = ""
                meta["detected_type"] = "unknown"

    extracted_text = _normalize_text(extracted_text)
    max_for_ai = _env_int("AI_ANALYSIS_MAX_CHARS", 20000)
    if len(extracted_text) > max_for_ai:
        extracted_text = _hierarchical_summarize(extracted_text, ai_model=ai_model, target_chars=_env_int("AI_ANALYSIS_TARGET_CHARS", 12000))

    prompt = (
        f"请分析下面的文档内容（file_id: {file_id}）。分析类型：{analysis_type}\n\n"
        "输出要求：只输出严格 JSON（不要代码块、不要说明文字）。\n"
        "1) structure：返回章节标题、层级、可能的编号、关键表格/附件提示；\n"
        "2) style：返回可观察到的排版风格（标题层级、列表、表格、引用等），不要凭空臆测字体颜色；\n"
        "3) content：返回 200-400 字摘要 + 10 条关键要点。\n\n"
        f"【内容】\n{extracted_text}"
    )

    ai_response = _strip_think(ai_client.generate_completion(prompt, model=ai_model))
    try:
        return _robust_json_loads(ai_response)
    except Exception:
        return {"raw_analysis": ai_response, "metadata": meta}

async def _extract_content_with_meta(file_id: str, format: str) -> tuple[str, dict]:
    try:
        file_content = await download_file_from_backend(file_id)
    except Exception as e:
        return (f"无法提取内容：文档 {file_id} 下载失败 - {str(e)}", {"error": str(e)})
    
    try:
        try:
            text, meta = _extract_pdf_text(file_content)
            if text:
                meta["detected_type"] = "pdf"
                return (text if format in {"markdown", "plain"} else text, meta)
        except Exception:
            pass
        
        try:
            text, meta = _extract_docx_text(file_content)
            if text:
                meta["detected_type"] = "docx"
                return (text if format in {"markdown", "plain"} else text, meta)
        except Exception:
            pass
        
        try:
            text = file_content.decode("utf-8", errors="ignore")
            if text.strip():
                return (_normalize_text(text), {"detected_type": "text"})
        except Exception:
            pass
        
        return (f"无法识别文件格式或文件内容为空", {"detected_type": "unknown"})
    except Exception as e:
        return (f"提取内容失败：{str(e)}", {"error": str(e)})

async def _extract_content_logic(file_id: str, format: str):
    text, _meta = await _extract_content_with_meta(file_id, format)
    text = _normalize_text(text)
    max_ret = _env_int("CONTENT_EXTRACTOR_MAX_RETURN_CHARS", 200000)
    head = _env_int("CONTENT_EXTRACTOR_HEAD_CHARS", 120000)
    tail = _env_int("CONTENT_EXTRACTOR_TAIL_CHARS", 40000)
    if max_ret > 0 and len(text) > max_ret and head > 0 and tail > 0:
        head_text = text[:head]
        tail_text = text[-tail:]
        text = _normalize_text(head_text + "\n\n...[TRUNCATED]...\n\n" + tail_text)
    return text

async def _match_template_logic(content_file_ids: list[str], template_file_id: str, keep_styles: bool, ai_model: str | None = None):
    content_summaries: list[dict] = []
    for fid in content_file_ids[:10]:
        try:
            text, meta = await _extract_content_with_meta(fid, "markdown")
            text = _normalize_text(text)
            if len(text) > _env_int("AI_MATCHER_MAX_CHARS_PER_FILE", 12000):
                text = _hierarchical_summarize(text, ai_model=ai_model, target_chars=_env_int("AI_MATCHER_TARGET_CHARS_PER_FILE", 6000))
            content_summaries.append({"file_id": fid, "summary": text, "metadata": meta})
        except Exception as e:
            content_summaries.append({"file_id": fid, "error": str(e)})

    template_summary = ""
    template_meta: dict = {}
    if template_file_id and template_file_id != "none":
        try:
            t, template_meta = await _extract_content_with_meta(template_file_id, "markdown")
            t = _normalize_text(t)
            if len(t) > _env_int("AI_MATCHER_MAX_CHARS_TEMPLATE", 16000):
                t = _hierarchical_summarize(t, ai_model=ai_model, target_chars=_env_int("AI_MATCHER_TARGET_CHARS_TEMPLATE", 8000))
            template_summary = t
        except Exception as e:
            template_summary = ""
            template_meta = {"error": str(e)}

    prompt = (
        "你是文档合并与排版策略助手。\n"
        f"目标：将多个内容文件合并到模板中，保留样式：{keep_styles}。\n"
        "请输出严格 JSON（不要代码块、不要说明文字），包含：\n"
        '1) "sections_mapping": 章节/条款映射；2) "tables": 表格如何迁移；3) "notes": 风险与缺失字段；4) "order": 推荐章节顺序。\n\n'
        f"【模板摘要】\n{template_summary}\n\n【内容摘要】\n{json.dumps(content_summaries, ensure_ascii=False)}"
    )
    
    ai_response = ai_client.generate_completion(prompt, model=ai_model)
    try:
        return _robust_json_loads(ai_response)
    except:
        return {"plan": ai_response, "template_metadata": template_meta}

async def _generate_document_logic(content: str, template_file_id: str, output_format: str, preset_template: str | None = None, ai_model: str | None = None) -> dict:
    template_type = 'default'
    
    if preset_template:
        preset_lower = preset_template.lower()
        if 'resume' in preset_lower or '简历' in preset_template:
            template_type = 'resume'
        elif 'report' in preset_lower or '报告' in preset_template:
            template_type = 'report'
        elif 'contract' in preset_lower or '合同' in preset_template:
            template_type = 'contract'
        elif 'meeting' in preset_lower or '会议' in preset_template:
            template_type = 'meeting'
        elif 'proposal' in preset_lower or '提案' in preset_template:
            template_type = 'proposal'
        elif 'invoice' in preset_lower or '发票' in preset_template:
            template_type = 'invoice'
    
    print(f"Generating document: template_type={template_type}, content_length={len(content) if content else 0}")
    
    try:
        doc_content = create_document_from_content(content, template_type, ai_model=ai_model)
        
        import time
        timestamp = int(time.time())
        filename = f"generated-{template_type}-{timestamp}.docx"
        file_id = await upload_generated_document(doc_content, filename)
        
        return {"result_file_id": file_id, "filename": filename, "template_type": template_type}
    except Exception as e:
        import traceback
        traceback.print_exc()
        error_msg = f"Failed to generate/upload document: {str(e)}"
        print(error_msg)
        return {"error": error_msg}

async def _modify_document_logic(file_id: str, modifications: str, ai_model: str | None = None) -> dict:
    file_content = await download_file_from_backend(file_id)
    
    if not file_content:
        return {"error": "Failed to download document for modification"}
    
    try:
        modified_content = modify_document_with_content(file_content, modifications, ai_model=ai_model)
        
        import time
        timestamp = int(time.time())
        filename = f"modified-{timestamp}.docx"
        new_file_id = await upload_generated_document(modified_content, filename)
        
        return {"result_file_id": new_file_id, "filename": filename}
    except Exception as e:
        error_msg = f"Failed to modify/upload document: {str(e)}"
        print(error_msg)
        return {"error": error_msg}

async def _structured_extractor_logic(file_id: str, template_type: str | None = None, ai_model: str | None = None) -> dict:
    text, meta = await _extract_content_with_meta(file_id, "markdown")
    text = _normalize_text(text)

    inferred = template_type or "auto"
    if inferred == "auto":
        inferred = _infer_template_type(text[:20000])

    parsed = parse_ai_content_for_template(text, inferred, ai_model=ai_model)
    return {"template_type": inferred, "data": parsed, "metadata": meta, "content_length": len(text)}


# ==================== 文档审查工具 ====================

async def _document_reviewer_logic(file_id: str, review_type: str = "general", ai_model: str | None = None) -> dict:
    """审查文档，识别风险点并生成批注"""
    text, meta = await _extract_content_with_meta(file_id, "markdown")
    text = _normalize_text(text)
    
    if not text:
        return {"error": "无法提取文档内容", "annotations": [], "summary": "", "risk_level": "unknown"}
    
    # 根据审查类型选择不同的 prompt
    review_prompts = {
        "legal": """你是一位资深法务专家。请审查以下文档，识别潜在的法律风险和问题。

审查要点：
1. 合同条款是否完整（主体、标的、期限、违约责任等）
2. 是否存在不公平条款或霸王条款
3. 权利义务是否对等
4. 争议解决条款是否合理
5. 是否存在法律漏洞

请以JSON格式返回审查结果：
{
    "annotations": [
        {
            "position": "第X段/第X条",
            "original_text": "原文片段",
            "issue_type": "风险类型",
            "severity": "high/medium/low",
            "comment": "问题描述和修改建议"
        }
    ],
    "summary": "整体审查总结（200字以内）",
    "risk_level": "critical/high/medium/low",
    "recommendations": ["建议1", "建议2"]
}

文档内容：
""",
        "compliance": """你是一位合规审查专家。请审查以下文档的合规性。

审查要点：
1. 是否符合相关法规要求
2. 是否包含必要的声明和披露
3. 数据隐私和安全条款
4. 知识产权相关条款
5. 行业特定合规要求

请以JSON格式返回审查结果：
{
    "annotations": [
        {
            "position": "第X段/第X条",
            "original_text": "原文片段",
            "issue_type": "合规问题类型",
            "severity": "high/medium/low",
            "comment": "问题描述和整改建议"
        }
    ],
    "summary": "合规审查总结（200字以内）",
    "risk_level": "critical/high/medium/low",
    "compliance_gaps": ["缺失项1", "缺失项2"]
}

文档内容：
""",
        "risk": """你是一位风险评估专家。请审查以下文档中的潜在风险。

审查要点：
1. 财务风险（付款条款、违约金等）
2. 运营风险（交付时间、质量标准等）
3. 声誉风险（保密条款、竞业禁止等）
4. 技术风险（知识产权、技术转让等）
5. 不可抗力和意外情况

请以JSON格式返回审查结果：
{
    "annotations": [
        {
            "position": "第X段/第X条",
            "original_text": "原文片段",
            "risk_category": "财务/运营/声誉/技术/其他",
            "severity": "high/medium/low",
            "probability": "高/中/低",
            "impact": "影响描述",
            "mitigation": "风险缓解建议"
        }
    ],
    "summary": "风险评估总结（200字以内）",
    "risk_level": "critical/high/medium/low",
    "risk_matrix": {"high_high": 0, "high_medium": 0, "medium_medium": 0, "low": 0}
}

文档内容：
""",
        "general": """你是一位文档审查专家。请审查以下文档，识别问题和改进建议。

审查要点：
1. 内容完整性
2. 逻辑一致性
3. 表述清晰度
4. 潜在歧义或模糊之处
5. 格式和结构问题

请以JSON格式返回审查结果：
{
    "annotations": [
        {
            "position": "第X段",
            "original_text": "原文片段",
            "issue_type": "问题类型",
            "severity": "high/medium/low",
            "comment": "问题描述和修改建议"
        }
    ],
    "summary": "审查总结（200字以内）",
    "risk_level": "high/medium/low",
    "improvements": ["改进建议1", "改进建议2"]
}

文档内容：
"""
    }
    
    prompt = review_prompts.get(review_type, review_prompts["general"])
    
    # 如果文档太长，进行摘要
    max_chars = _env_int("AI_REVIEW_MAX_CHARS", 20000)
    if len(text) > max_chars:
        text = _hierarchical_summarize(text, ai_model=ai_model, target_chars=max_chars)
    
    full_prompt = prompt + text + "\n\n要求：只输出严格的 JSON（不要代码块、不要说明文字）。"
    
    ai_response = _strip_think(ai_client.generate_completion(full_prompt, model=ai_model))
    
    try:
        result = _robust_json_loads(ai_response)
        result["review_type"] = review_type
        result["document_length"] = len(text)
        return result
    except Exception as e:
        return {
            "error": f"解析审查结果失败: {str(e)}",
            "raw_response": ai_response[:2000],
            "annotations": [],
            "summary": "审查完成但结果解析失败",
            "risk_level": "unknown"
        }


# ==================== AI 处理器工具 ====================

async def _ai_processor_logic(prompt: str, ai_model: str | None = None) -> dict:
    """通用 AI 处理器，用于工作流中的自定义处理"""
    if not prompt:
        return {"error": "Prompt is required", "content": ""}
    
    ai_response = _strip_think(ai_client.generate_completion(prompt, model=ai_model))
    
    return {
        "content": ai_response,
        "prompt_length": len(prompt),
        "response_length": len(ai_response)
    }


# ==================== 音频转录工具 ====================

async def _audio_transcriber_logic(file_id: str, generate_minutes: bool = True, ai_model: str | None = None) -> dict:
    """转录音频并可选生成会议纪要"""
    try:
        file_content = await download_file_from_backend(file_id)
    except Exception as e:
        return {"error": f"下载音频文件失败: {str(e)}", "transcript": ""}
    
    # 尝试使用 Whisper 或其他 ASR 服务进行转录
    # 这里我们提供一个模拟实现，实际使用时需要集成真正的 ASR 服务
    transcript = ""
    speakers = []
    
    try:
        # 尝试使用 OpenAI Whisper API 或本地 Whisper
        # 如果配置了 WHISPER_API_URL，使用远程服务
        whisper_url = os.getenv("WHISPER_API_URL", "")
        
        if whisper_url:
            async with httpx.AsyncClient(timeout=300.0) as client:
                files = {'file': ('audio.mp3', io.BytesIO(file_content), 'audio/mpeg')}
                resp = await client.post(
                    whisper_url,
                    files=files,
                    headers={"Authorization": f"Bearer {AI_API_KEY}"} if AI_API_KEY else {}
                )
                if resp.status_code == 200:
                    result = resp.json()
                    transcript = result.get("text", "")
                    speakers = result.get("segments", [])
        else:
            # 如果没有配置 Whisper，使用 AI 进行模拟（仅用于演示）
            # 实际生产环境应该集成真正的 ASR 服务
            transcript = "[音频转录功能需要配置 WHISPER_API_URL 环境变量。请配置 OpenAI Whisper API 或兼容的 ASR 服务。]"
            
    except Exception as e:
        print(f"ASR Error: {e}")
        transcript = f"[转录失败: {str(e)}]"
    
    result = {
        "transcript": transcript,
        "speakers": speakers,
        "audio_file_id": file_id
    }
    
    # 如果需要生成会议纪要
    if generate_minutes and transcript and not transcript.startswith("["):
        minutes_prompt = f"""请根据以下会议转录内容，生成一份结构化的会议纪要。

转录内容：
{transcript}

请以JSON格式返回：
{{
    "title": "会议主题",
    "date": "会议时间（如能识别）",
    "attendees": ["参会人员列表"],
    "topics": ["议题1", "议题2"],
    "discussion": "主要讨论内容摘要",
    "decisions": ["决议1", "决议2"],
    "action_items": [
        {{"item": "待办事项", "owner": "负责人", "deadline": "截止时间"}}
    ],
    "summary": "会议总结（200字以内）"
}}

要求：只输出严格的 JSON。"""

        ai_response = _strip_think(ai_client.generate_completion(minutes_prompt, model=ai_model))
        
        try:
            minutes_data = _robust_json_loads(ai_response)
            result["summary"] = minutes_data.get("summary", "")
            result["action_items"] = minutes_data.get("action_items", [])
            result["minutes_data"] = minutes_data
            
            # 生成会议纪要文档
            doc_content = create_document_from_content(
                json.dumps(minutes_data, ensure_ascii=False, indent=2),
                "meeting",
                ai_model=ai_model
            )
            
            import time
            timestamp = int(time.time())
            filename = f"meeting-minutes-{timestamp}.docx"
            result_file_id = await upload_generated_document(doc_content, filename)
            result["result_file_id"] = result_file_id
            
        except Exception as e:
            result["minutes_error"] = str(e)
    
    return result

@mcp.tool()
async def document_analyzer(file_id: str, analysis_type: str = "structure", ai_model: str | None = None) -> str:
    """分析文档结构、内容类型或样式。"""
    result = await _analyze_document_logic(file_id, analysis_type, ai_model=ai_model)
    return json.dumps(result, indent=2, ensure_ascii=False)

@mcp.tool()
async def content_extractor(file_id: str, format: str = "markdown") -> str:
    """以指定格式从文档中提取内容。"""
    result = await _extract_content_logic(file_id, format)
    return result

@mcp.tool()
async def template_matcher(content_file_ids: list[str], template_file_id: str, keep_styles: bool = True, ai_model: str | None = None) -> str:
    """将内容与模板匹配并生成布局计划。"""
    result = await _match_template_logic(content_file_ids, template_file_id, keep_styles, ai_model=ai_model)
    return json.dumps(result, indent=2, ensure_ascii=False)

@mcp.tool()
async def document_generator(content: str, template_file_id: str, output_format: str = "docx", preset_template: str | None = None, ai_model: str | None = None) -> str:
    """生成最终文档。"""
    result = await _generate_document_logic(content, template_file_id, output_format, preset_template, ai_model=ai_model)
    return json.dumps(result, indent=2, ensure_ascii=False)

@mcp.tool()
async def document_modifier(file_id: str, modifications: str, ai_model: str | None = None) -> str:
    """根据修改要求修改现有文档。"""
    result = await _modify_document_logic(file_id, modifications, ai_model=ai_model)
    return json.dumps(result, indent=2, ensure_ascii=False)

@mcp.tool()
async def structured_extractor(file_id: str, template_type: str = "auto", ai_model: str | None = None) -> str:
    """对文档做结构化信息抽取。"""
    result = await _structured_extractor_logic(file_id, template_type, ai_model=ai_model)
    return json.dumps(result, indent=2, ensure_ascii=False)

@mcp.tool()
async def document_reviewer(file_id: str, review_type: str = "general", ai_model: str | None = None) -> str:
    """审查文档，识别风险点并生成批注。review_type: legal/compliance/risk/general"""
    result = await _document_reviewer_logic(file_id, review_type, ai_model=ai_model)
    return json.dumps(result, indent=2, ensure_ascii=False)

@mcp.tool()
async def ai_processor(prompt: str, ai_model: str | None = None) -> str:
    """通用AI处理器，用于自定义处理任务。"""
    result = await _ai_processor_logic(prompt, ai_model=ai_model)
    return json.dumps(result, indent=2, ensure_ascii=False)

@mcp.tool()
async def audio_transcriber(file_id: str, generate_minutes: bool = True, ai_model: str | None = None) -> str:
    """转录音频文件并可选生成会议纪要。"""
    result = await _audio_transcriber_logic(file_id, generate_minutes, ai_model=ai_model)
    return json.dumps(result, indent=2, ensure_ascii=False)

from fastapi import FastAPI
from mcp.server import Server
from mcp.server.sse import SseServerTransport
from mcp.types import Tool, TextContent

app = FastAPI()
mcp_server = Server("docai-mcp")

@mcp_server.list_tools()
async def list_tools() -> list[Tool]:
    return [
        Tool(name="document_analyzer", description="分析文档结构、内容或样式", inputSchema={
            "type": "object", 
            "properties": {
                "file_id": {"type": "string"}, 
                "analysis_type": {"type": "string", "enum": ["structure", "style", "content"]}
            }
        }),
        Tool(name="content_extractor", description="提取文档内容", inputSchema={
            "type": "object", 
            "properties": {
                "file_id": {"type": "string"}, 
                "format": {"type": "string", "enum": ["markdown", "plain"]}
            }
        }),
        Tool(name="template_matcher", description="匹配模板并生成布局计划", inputSchema={
            "type": "object", 
            "properties": {
                "content_file_ids": {"type": "array", "items": {"type": "string"}}, 
                "template_file_id": {"type": "string"},
                "keep_styles": {"type": "boolean"}
            }
        }),
        Tool(name="document_generator", description="生成文档", inputSchema={
            "type": "object", 
            "properties": {
                "content": {"type": "string"}, 
                "template_file_id": {"type": "string"},
                "output_format": {"type": "string"},
                "preset_template": {"type": "string"}
            }
        }),
        Tool(name="document_modifier", description="修改现有文档", inputSchema={
            "type": "object",
            "properties": {
                "file_id": {"type": "string"},
                "modifications": {"type": "string"}
            }
        }),
        Tool(name="structured_extractor", description="结构化信息抽取", inputSchema={
            "type": "object",
            "properties": {
                "file_id": {"type": "string"},
                "template_type": {"type": "string"},
                "ai_model": {"type": "string"}
            }
        }),
        Tool(name="document_reviewer", description="审查文档，识别风险点并生成批注", inputSchema={
            "type": "object",
            "properties": {
                "file_id": {"type": "string"},
                "review_type": {"type": "string", "enum": ["legal", "compliance", "risk", "general"]},
                "ai_model": {"type": "string"}
            }
        }),
        Tool(name="ai_processor", description="通用AI处理器", inputSchema={
            "type": "object",
            "properties": {
                "prompt": {"type": "string"},
                "ai_model": {"type": "string"}
            }
        }),
        Tool(name="audio_transcriber", description="转录音频并生成会议纪要", inputSchema={
            "type": "object",
            "properties": {
                "file_id": {"type": "string"},
                "generate_minutes": {"type": "boolean"},
                "ai_model": {"type": "string"}
            }
        }),
    ]

@mcp_server.call_tool()
async def call_tool(name: str, arguments: dict) -> list[TextContent]:
    if name == "document_analyzer":
        res = await _analyze_document_logic(arguments["file_id"], arguments.get("analysis_type", "structure"))
        return [TextContent(type="text", text=json.dumps(res, indent=2, ensure_ascii=False))]
    elif name == "content_extractor":
        res = await _extract_content_logic(arguments["file_id"], arguments.get("format", "markdown"))
        return [TextContent(type="text", text=res)]
    elif name == "template_matcher":
        res = await _match_template_logic(arguments["content_file_ids"], arguments["template_file_id"], arguments.get("keep_styles", True))
        return [TextContent(type="text", text=json.dumps(res, indent=2, ensure_ascii=False))]
    elif name == "document_generator":
        res = await _generate_document_logic(arguments["content"], arguments["template_file_id"], arguments.get("output_format", "docx"), arguments.get("preset_template"))
        return [TextContent(type="text", text=json.dumps(res, indent=2, ensure_ascii=False))]
    elif name == "document_modifier":
        res = await _modify_document_logic(arguments["file_id"], arguments["modifications"])
        return [TextContent(type="text", text=json.dumps(res, indent=2, ensure_ascii=False))]
    elif name == "structured_extractor":
        res = await _structured_extractor_logic(arguments["file_id"], arguments.get("template_type", "auto"), ai_model=arguments.get("ai_model"))
        return [TextContent(type="text", text=json.dumps(res, indent=2, ensure_ascii=False))]
    elif name == "document_reviewer":
        res = await _document_reviewer_logic(arguments["file_id"], arguments.get("review_type", "general"), ai_model=arguments.get("ai_model"))
        return [TextContent(type="text", text=json.dumps(res, indent=2, ensure_ascii=False))]
    elif name == "ai_processor":
        res = await _ai_processor_logic(arguments["prompt"], ai_model=arguments.get("ai_model"))
        return [TextContent(type="text", text=json.dumps(res, indent=2, ensure_ascii=False))]
    elif name == "audio_transcriber":
        res = await _audio_transcriber_logic(arguments["file_id"], arguments.get("generate_minutes", True), ai_model=arguments.get("ai_model"))
        return [TextContent(type="text", text=json.dumps(res, indent=2, ensure_ascii=False))]
    return []

@app.post("/api/tools/{name}/invoke")
async def invoke_tool(name: str, arguments: dict):
    ai_model = arguments.get("ai_model")
    if name == "document_analyzer":
        return await _analyze_document_logic(arguments["file_id"], arguments.get("analysis_type", "structure"), ai_model=ai_model)
    elif name == "content_extractor":
        content, meta = await _extract_content_with_meta(arguments["file_id"], arguments.get("format", "markdown"))
        content = _normalize_text(content)
        max_ret = _env_int("CONTENT_EXTRACTOR_MAX_RETURN_CHARS", 200000)
        head = _env_int("CONTENT_EXTRACTOR_HEAD_CHARS", 120000)
        tail = _env_int("CONTENT_EXTRACTOR_TAIL_CHARS", 40000)
        truncated = False
        if max_ret > 0 and len(content) > max_ret and head > 0 and tail > 0:
            content = _normalize_text(content[:head] + "\n\n...[TRUNCATED]...\n\n" + content[-tail:])
            truncated = True
        meta = dict(meta or {})
        meta["content_length"] = len(content)
        meta["truncated"] = truncated
        return {"content": content, "metadata": meta}
    elif name == "template_matcher":
        return await _match_template_logic(arguments["content_file_ids"], arguments["template_file_id"], arguments.get("keep_styles", True), ai_model=ai_model)
    elif name == "document_generator":
        return await _generate_document_logic(arguments["content"], arguments["template_file_id"], arguments.get("output_format", "docx"), arguments.get("preset_template"), ai_model=ai_model)
    elif name == "document_modifier":
        return await _modify_document_logic(arguments["file_id"], arguments["modifications"], ai_model=ai_model)
    elif name == "structured_extractor":
        return await _structured_extractor_logic(arguments["file_id"], arguments.get("template_type", "auto"), ai_model=ai_model)
    elif name == "document_reviewer":
        return await _document_reviewer_logic(arguments["file_id"], arguments.get("review_type", "general"), ai_model=ai_model)
    elif name == "ai_processor":
        return await _ai_processor_logic(arguments["prompt"], ai_model=ai_model)
    elif name == "audio_transcriber":
        return await _audio_transcriber_logic(arguments["file_id"], arguments.get("generate_minutes", True), ai_model=ai_model)
    raise HTTPException(status_code=404, detail="Tool not found")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=3000)
